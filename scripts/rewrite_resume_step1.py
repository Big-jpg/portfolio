from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(
    r"C:\Users\rossf\Downloads\Ross_Farrell_Vercel_Solutions_Architect_Resume_Rewritten.docx"
)
OUTPUT = Path(
    r"C:\Users\rossf\Desktop\portfolio\Ross_Farrell_Vercel_Solution_Architect_Resume_Review_01.docx"
)


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace text while retaining paragraph style and representative run formatting."""
    run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    paragraph._p.append(run)


def paragraph_with_exact_text(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph matching {text!r}; found {len(matches)}")
    return matches[0]


document = Document(SOURCE)

replace_paragraph_text(
    paragraph_with_exact_text(
        document, "SOLUTION ARCHITECT | MODERN WEB, AI, AND ENTERPRISE PLATFORMS"
    ),
    "SOLUTION ARCHITECT | MODERN WEB, AI, AND ENTERPRISE PLATFORMS",
)

replace_paragraph_text(
    paragraph_with_exact_text(
        document, "Portfolio: hypecoding.dev | Vercel contribution: vercel/examples #594"
    ),
    "Portfolio: hypecoding.dev | GitHub: github.com/Big-jpg",
)

replace_paragraph_text(
    paragraph_with_exact_text(
        document,
        "Solution Architect with 18 years of experience delivering enterprise systems across regulated and operational environments, including six years focused on cloud architecture, modern web applications, AI, automation, and data platforms. Operates across executive discovery, solution design, production engineering, and hands-on implementation. Turns fragmented technology landscapes into coherent systems, delivery standards, and operating models.",
    ),
    "Customer-facing Solution Architect with six years of experience leading technical discovery, solution design, implementation, and post-go-live support for enterprise clients, backed by 18 years across regulated and operational environments. Hands-on across TypeScript, React, Next.js, cloud, AI, automation, data platforms, integration, and identity. Translates executive priorities and complex technical constraints into secure production systems, measurable outcomes, and reusable delivery patterns.",
)

replace_paragraph_text(
    paragraph_with_exact_text(document, "SELECTED HIGHLIGHT"),
    "SELECTED ARCHITECTURE WORK",
)

highlight = paragraph_with_exact_text(
    document,
    "Early Vercel AI adopter: Identified and documented an early GPT-4 compatibility issue in Vercel's original Next.js chatbot example (vercel/examples #594), helping later developers understand and resolve the request-payload failure.",
)
highlight.style = document.styles["Resume Bullet"]
replace_paragraph_text(
    highlight,
    "• GitHub Motion Graph: Built and deployed a Next.js 16 application using App Router, React Server Components, Edge and Node API routes, Neon and Drizzle, and Vercel Analytics to visualise GitHub development activity.",
)

core_capabilities_heading = paragraph_with_exact_text(document, "CORE CAPABILITIES")
core_capabilities_heading.insert_paragraph_before(
    "• InvoicePipe: Architected a tenant-isolated invoice-processing platform combining Next.js, Azure AI extraction, a containerised Python service, PostgreSQL, typed API responses, input validation, correlation IDs, and structured logging.",
    style="Resume Bullet",
)
core_capabilities_heading.insert_paragraph_before(
    "• ModelViz: Built a Next.js and TypeScript application that parses Power BI and Microsoft Fabric semantic models locally, classifies data-source modes, inspects row-level security, and renders interactive architecture diagrams.",
    style="Resume Bullet",
)

capability_rows = [
    (
        "Customer lifecycle",
        "Technical discovery, solution qualification, target architecture, workshops, proofs of concept, migration strategy, executive communication, and post-sale adoption",
    ),
    (
        "Modern web and AI",
        "TypeScript, JavaScript, React, Next.js, Node.js, API and edge design, SSR/SSG, Vercel, OpenAI APIs, and AI-powered applications",
    ),
    (
        "Enterprise platforms",
        "Azure, AWS, Microsoft Fabric, integration architecture, SQL, Python, PySpark, Docker, Kubernetes, observability, performance, and cost optimisation",
    ),
    (
        "Security and governance",
        "SAML, SSO, LDAP, RBAC, RLS, networking, audit telemetry, data protection, lineage, CI/CD, and regulated delivery",
    ),
]

table = document.tables[0]
if len(table.rows) != len(capability_rows):
    raise ValueError("Unexpected core-capabilities table shape")

for row, (label, detail) in zip(table.rows, capability_rows):
    replace_paragraph_text(row.cells[0].paragraphs[0], label)
    replace_paragraph_text(row.cells[1].paragraphs[0], detail)

document.save(OUTPUT)
print(OUTPUT)
